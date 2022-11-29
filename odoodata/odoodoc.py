import datetime as dt
import os
from pprint import pprint

import pandas as pd
import glob
import docx
from docx2pdf import convert
from docx.enum.text import WD_BREAK
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from constants import *
import odoodata.odoodata as od
from utils import *
import gdrive_data.auth as auth
import gdrive_data.main as main


def format_datetime(date=None, time=None):
    if date is None and time is None:
        start = dt.date.today()
    elif date is not None:
        date = date.split('-')
        date = [int(value) for value in date]
        if time is not None:
            time = time.split(':')
            time = [int(value) for value in time]
            start = dt.datetime(date[0],date[1],date[2],time[0]+1,time[1],time[0])
        else:
            start = dt.date(date[0],date[1],date[2])
    return start


class Doc():
    def __init__(self, start_date=None, end_date=None, 
                start_time=None, end_time=None, h_font='Alef'
            ):

        self.path = f'{os.environ["USERPROFILE"]}/Documents/'
        self.path_to_catalog = f'{os.environ["USERPROFILE"]}/Documents/{CATALOG_FOLDER}/'

        self.comp = od.getData(PASSWORD, EMAIL, DOMAIN, start_date, end_date, start_time, end_time)
        self.h_font = h_font
        self.doc = docx.Document()

        if start_date is None and end_date is None:
            self.start = format_datetime()
            self.date = f'{self.start:%d %b %y}'
        elif start_date is None and end_date != None:
            self.start = format_datetime(date=end_date)
            self.date = f'{self.start:%d %b %y}'
        elif start_date != None and end_date is None:
            self.start = format_datetime(date=start_date)
            self.date = f'{self.start:%d %b %y}'
        elif start_date != None and end_date != None:
            self.start = format_datetime(date=start_date, time=start_time)
            end_dt = format_datetime(date=end_date, time=end_time)
            self.date = f'{self.start:%d %b %y_%I-%M %p} - {end_dt:%d %b %y_%I-%M %p}'
    def docHeader():
        def decorator(func):
            def wrapper(self, add_pic, *args, **kwargs):
                print('creating document...')
                section = self.doc.sections[0]
                header = section.header
                header_para = header.paragraphs[0]
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = header_para.add_run()
                font = run.font
                font.name = self.h_font
                font.size = Pt(15)
                run.text = COMPANY
                para = header.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.text = self.date
                func(self, add_pic, *args, **kwargs)
            return wrapper
        return decorator

    @docHeader()
    def createDoc(self, add_pic, dicts, t_title=[], product=[], images=[],
        pro_categ=[], t_style='Table Grid', f_name='Alef',
        ):

        def get_image_hyperlink(img_path):
            hyp_name = 'file:///' + img_path
            return hyp_name

        def add_hyperlink(paragraph, text, url):
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
            hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

            new_run = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')

            new_run.append(rPr)
            new_run.text = text
            hyperlink.append(new_run)

            r = paragraph.add_run ()
            r._r.append (hyperlink)

            r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
            r.font.underline = True
            r.font.size = Pt(7)
            return hyperlink

        def find_image(img, img_v, category):
            images = glob.glob(f'{self.path_to_catalog}{category}/{img_v}*')
            if len(images)==0:
                images = glob.glob(f'{self.path_to_catalog}{category}/{img}*')
            if len(images)==0:
                images = glob.glob(f'{self.path_to_catalog}**/{img}*', recursive=True) #check all the folders for any matching file that bears the name of the product. Including all variants.
                if len(images) > 0:
                    img_path = None
                    for im in images:
                        file_path, file_ext = os.path.splitext(im)
                        file_name = file_path.split('\\')
                        try:
                            os.rename(im, f'{self.path_to_catalog}{category}/{file_name[len(file_name)-1]}{file_ext}') # move all the found files to their appropriate folers
                        except: # if the file cannot be moved for some reason, quickly check if this is the image we're looking for and set the return value before moving on.
                            if img_v in file_ext:
                                img_path = im
                    if not img_path: # if img_path was not set during the loop, check the appropriate folders again for the variants and the product respectively
                        images = glob.glob(f'{self.path_to_catalog}{category}/{img_v}*')
                        if len(images) == 0:
                            images = glob.glob(f'{self.path_to_catalog}{category}/{img}*')
                    else:
                        return img_path
            return images[0]
        try:    
            os.chdir('gdrive_data\\')
            categs = set()
            [categs.add(i) for i in pro_categ]
            aut = auth.auth(SCOPES)
            os.chdir('..')
            creds = aut.getCred()
            service = main.build('drive', 'v3', credentials=creds)
            drive_pics = main.search_file(service, categs)
        except:
            drive_pics = None
        
        if len(dicts) != len(t_title):
            raise ValueError('Length of table and table title must match')
        if len(dicts) != len(add_pic):
            raise ValueError('Length of bool list must match table length')

        for n, table in enumerate(dicts):
            self.doc.add_heading(t_title[n], 2)
            if add_pic[n] == True:
                if len(product) > len(table):
                    raise ValueError('Length of table must be == product length')
                if len(images) > len(table):
                    raise ValueError('Length of table must be == images length')
                if len(pro_categ) > len(table):
                    raise ValueError('Length of table must be == product category length')
                t = self.doc.add_table(len(table)+1, len(table[0].keys())+1)
                t.cell(0,len(table[0].keys())).text = 'Image'
            else:
                t = self.doc.add_table(len(table)+1, len(table[0].keys()))
            
            t.style = t_style
            font = t.style.font
            font.name = f_name

            j=0
            for key in table[0].keys():
                t.cell(0,j).text = key
                j+=1
            i=1
            for record in table:
                j=0
                for value in record.values():
                    if isinstance(value, int) or isinstance(value, float):
                        value = f'{value:,.2f}'
                    t.cell(i,j).text = value
                    t.cell(i,j).paragraphs[0].runs[0].font.size = Pt(12)
                    j+=1
                i+=1
            i=1
            if add_pic[n] == True:
                for record in range(len(images)):
                    try:
                        cell = t.cell(i,len(table[0].keys()))
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        img_path = find_image(images[i-1], product[i-1], pro_categ[i-1])
                        run.add_picture(img_path, width = Inches(.7), height = Inches(.8))
                        run.add_break(WD_BREAK.LINE)
                        if drive_pics is not None:
                            try:
                                link = drive_pics[pro_categ[i-1]][product[i-1]]
                            except KeyError:
                                link = drive_pics[pro_categ[i-1]][images[i-1]]
                            except:
                                link = get_image_hyperlink(img_path)
                        else:
                            link = get_image_hyperlink(img_path)
                        add_hyperlink(paragraph, 'open', link)
                    except:
                        t.cell(i,len(table[0].keys())).text = 'no image'
                    i+=1
        
    def createPath(self, report_name):
        def create_file(new_month_dir):
            #keep trying to delete doc and set status==1(a good thing) to proceed
            doc_in_use = True
            i=1
            while doc_in_use:
                if os.path.exists(f"{new_month_dir}/{report_name} ({i}).docx") or os.path.exists(f"{new_month_dir}/{report_name} ({i}).pdf"):
                    try:
                        os.remove(f"{new_month_dir}/{report_name} ({i}).docx")
                        status=1
                    except FileNotFoundError:
                        status=1
                        pass
                    except:
                        status=0
                        i+=1
                    if status==1:
                        try:
                            os.remove(f"{new_month_dir}/{report_name} ({i}).pdf")
                        except:
                            i+=1
                    
                else:
                    doc_in_use = False
            return i

        pro_dir = rf'{self.path}/{FOLDER_NAME}'
        try:
            new_month_dir = rf'{self.path}/{FOLDER_NAME}/{self.start:%b %y}'
        except:
            new_month_dir = rf'{self.path}/{FOLDER_NAME}/{self.date:%b %y}'
        if not os.path.exists(pro_dir):
            os.mkdir(pro_dir)
        if not os.path.exists(new_month_dir):
            os.mkdir(new_month_dir)
        i=0
        if os.path.exists(f"{new_month_dir}/{report_name}.docx") or os.path.exists(f"{new_month_dir}/{report_name}.pdf"):
            try:
                os.remove(f"{new_month_dir}/{report_name}.docx")
                status=1 # 1 means its a good thing and is expected
            except FileNotFoundError:
                status=1
                pass
            except:
                status=0 # 0 means not what was expected. Maybe the doc to be deleted is open or cannot delete for some reason
                i=create_file(new_month_dir) # so enter a loop and keep trying with a different name until status==1
            if status==1:
                try:
                    os.remove(f"{new_month_dir}/{report_name}.pdf")
                except:
                    i=create_file(new_month_dir)
        if i == 0:
            report_path = f'{new_month_dir}/{report_name}'
        else:
            report_path = f'{new_month_dir}/{report_name} ({i})'

        print('saving and converting doc..')
        self.doc.save(f'{report_path}.docx')
        convert(f'{report_path}.docx')
        try:
            os.remove(f'{report_path}.docx')
            os.startfile(f'{report_path}.pdf')
        except:
            pass
